"""RAG local híbrido para o MVP: BM25 + TF-IDF, com índice persistente."""
from __future__ import annotations

import datetime as dt
import hashlib
import json
import math
import re
import threading
from collections import Counter
from pathlib import Path
from typing import Any, List

import numpy as np
from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from config import settings
from errors import DocumentProcessingError, DocumentSecurityError
from retrieval.base import Retriever
from schemas.artifacts import TrechoRecuperado

_TOKEN_RE = re.compile(r"[\wÀ-ÿ-]+", flags=re.UNICODE)
_LOCK = threading.RLock()


def _tokens(text: str) -> list[str]:
    return [token.lower() for token in _TOKEN_RE.findall(text)]


def _normalize(values: np.ndarray) -> np.ndarray:
    if len(values) == 0:
        return values
    min_value, max_value = float(values.min()), float(values.max())
    if math.isclose(min_value, max_value):
        return np.ones_like(values) if max_value > 0 else np.zeros_like(values)
    return (values - min_value) / (max_value - min_value)


def _date_from_metadata(value: Any) -> dt.date | None:
    if not value:
        return None
    try:
        return dt.date.fromisoformat(str(value)[:10])
    except ValueError:
        return None


class LocalHybridRetriever(Retriever):
    def __init__(self, index_dir: str | None = None):
        self.index_dir = Path(index_dir or settings.RETRIEVAL_INDEX_DIR)
        self.index_dir.mkdir(parents=True, exist_ok=True)
        self.index_file = self.index_dir / "index.json"

    def _load(self) -> dict[str, Any]:
        if not self.index_file.exists():
            return {"documents": {}, "chunks": []}
        try:
            index = json.loads(self.index_file.read_text(encoding="utf-8"))
            index.setdefault("documents", {})
            index.setdefault("chunks", [])
            return index
        except (json.JSONDecodeError, OSError) as exc:
            raise DocumentProcessingError(
                "O índice local está corrompido. Restaure o backup ou recrie a base."
            ) from exc

    def _save(self, index: dict[str, Any]) -> None:
        temp = self.index_file.with_suffix(".tmp")
        temp.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
        temp.replace(self.index_file)

    def _extract_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                return "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
            if suffix == ".docx":
                doc = Document(path)
                paragraphs = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        paragraphs.append(" | ".join(cell.text for cell in row.cells))
                return "\n".join(paragraphs)
            if suffix == ".txt":
                return path.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            raise DocumentProcessingError(
                f"Não foi possível extrair o conteúdo de '{path.name}'."
            ) from exc
        raise DocumentProcessingError(f"Formato não suportado: {suffix}")

    @staticmethod
    def _chunk(text: str, max_chars: int = 1400, overlap_chars: int = 180) -> list[str]:
        paragraphs = [re.sub(r"\s+", " ", p).strip() for p in text.splitlines()]
        paragraphs = [p for p in paragraphs if p]
        chunks: list[str] = []
        current = ""
        for paragraph in paragraphs:
            candidate = f"{current} {paragraph}".strip()
            if len(candidate) <= max_chars:
                current = candidate
                continue
            if current:
                chunks.append(current)
            overlap = current[-overlap_chars:] if current else ""
            current = f"{overlap} {paragraph}".strip()
            while len(current) > max_chars:
                chunks.append(current[:max_chars])
                current = current[max_chars - overlap_chars :]
        if current:
            chunks.append(current)
        return [chunk for chunk in chunks if len(chunk) >= 80]

    def _guard_evaluation_content(self, filename: str, text: str) -> None:
        if not settings.ENABLE_KNOWLEDGE_BASE_GUARD:
            return
        normalized_name = filename.lower().replace(" ", "_").replace("-", "_")
        name_match = any(pattern in normalized_name for pattern in settings.suspicious_document_patterns)
        normalized_text = re.sub(r"\s+", " ", text.lower())
        content_markers = (
            "hipóteses esperadas",
            "resultado esperado da qualificação",
            "perguntas esperadas",
            "gabarito da demonstração",
            "guia de demonstração prospect-llm",
        )
        content_match = sum(marker in normalized_text for marker in content_markers) >= 2
        if name_match or content_match:
            raise DocumentSecurityError(
                "Este arquivo parece ser um guia, gabarito ou documento com resultados esperados. "
                "Mantenha-o fora da base de conhecimento para não contaminar a avaliação do RAG."
            )

    def ingest_document(
        self,
        file_path: str,
        classificacao: str,
        original_filename: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> str:
        path = Path(file_path)
        raw = path.read_bytes()
        document_id = hashlib.sha256(raw).hexdigest()
        filename = Path(original_filename or path.name).name
        text = self._extract_text(path)
        self._guard_evaluation_content(filename, text)
        if len(text.strip()) < 80:
            raise DocumentProcessingError(
                "O documento não possui texto suficiente para indexação. PDFs digitalizados exigem OCR."
            )
        chunks = self._chunk(text)
        if not chunks:
            raise DocumentProcessingError("Nenhum trecho útil foi obtido do documento.")

        metadata = dict(metadata or {})
        with _LOCK:
            index = self._load()
            created = dt.datetime.now(dt.timezone.utc).isoformat()
            current = index["documents"].get(document_id, {})
            current.update(
                {
                    "filename": filename,
                    "classificacao": classificacao,
                    "created_at": current.get("created_at", created),
                    "updated_at": created,
                    "chunk_count": len(chunks),
                    "active": bool(metadata.get("active", True)),
                    "use_in_generation": bool(metadata.get("use_in_generation", True)),
                    "version": metadata.get("version"),
                    "data_documento": metadata.get("data_documento"),
                    "valido_ate": metadata.get("valido_ate"),
                    "source_type": metadata.get("source_type"),
                }
            )
            index["documents"][document_id] = current
            index["chunks"] = [
                item for item in index["chunks"] if item.get("document_id") != document_id
            ]
            for position, chunk in enumerate(chunks):
                index["chunks"].append(
                    {
                        "id": f"{document_id}:{position}",
                        "document_id": document_id,
                        "filename": filename,
                        "classificacao": classificacao,
                        "text": chunk,
                    }
                )
            self._save(index)
        return document_id

    def list_documents(self) -> list[dict[str, Any]]:
        with _LOCK:
            index = self._load()
        return [dict(document_id=key, **value) for key, value in index["documents"].items()]

    def set_document_active(self, document_id: str, active: bool) -> None:
        with _LOCK:
            index = self._load()
            if document_id not in index["documents"]:
                raise DocumentProcessingError("Documento não encontrado no índice local.")
            index["documents"][document_id]["active"] = bool(active)
            index["documents"][document_id]["updated_at"] = dt.datetime.now(
                dt.timezone.utc
            ).isoformat()
            self._save(index)

    def remove_document(self, document_id: str) -> None:
        with _LOCK:
            index = self._load()
            index["documents"].pop(document_id, None)
            index["chunks"] = [
                item for item in index["chunks"] if item.get("document_id") != document_id
            ]
            self._save(index)

    def _allowed(self, chunk: dict[str, Any], documents: dict[str, Any]) -> bool:
        metadata = documents.get(chunk.get("document_id"), {})
        if not metadata.get("active", True) or not metadata.get("use_in_generation", True):
            return False
        valid_until = _date_from_metadata(metadata.get("valido_ate"))
        if valid_until and valid_until < dt.date.today():
            return False
        if not settings.is_cloud_llm or settings.ALLOW_RESTRICTED_CLOUD:
            return True
        return chunk.get("classificacao") in {"publico", "interno_nao_sensivel"}

    def _bm25_scores(self, query_tokens: list[str], docs_tokens: list[list[str]]) -> np.ndarray:
        if not docs_tokens:
            return np.array([], dtype=float)
        n_docs = len(docs_tokens)
        avgdl = sum(len(doc) for doc in docs_tokens) / max(n_docs, 1)
        dfs = Counter()
        for doc in docs_tokens:
            dfs.update(set(doc))
        scores = []
        k1, b = 1.5, 0.75
        for doc in docs_tokens:
            tf = Counter(doc)
            score = 0.0
            for term in query_tokens:
                df = dfs.get(term, 0)
                if df == 0:
                    continue
                idf = math.log(1 + (n_docs - df + 0.5) / (df + 0.5))
                freq = tf.get(term, 0)
                denom = freq + k1 * (1 - b + b * len(doc) / max(avgdl, 1))
                score += idf * (freq * (k1 + 1)) / max(denom, 1e-9)
            scores.append(score)
        return np.asarray(scores, dtype=float)

    def search(self, query: str, top_k: int = 8) -> List[TrechoRecuperado]:
        with _LOCK:
            index = self._load()
        documents = index["documents"]
        chunks = [chunk for chunk in index["chunks"] if self._allowed(chunk, documents)]
        if not chunks or not query.strip():
            return []

        texts = [chunk["text"] for chunk in chunks]
        query_tokens = _tokens(query)
        docs_tokens = [_tokens(text) for text in texts]
        lexical = self._bm25_scores(query_tokens, docs_tokens)
        try:
            vectorizer = TfidfVectorizer(
                lowercase=True,
                strip_accents="unicode",
                ngram_range=(1, 2),
                min_df=1,
                max_features=25000,
            )
            matrix = vectorizer.fit_transform(texts + [query])
            semantic = cosine_similarity(matrix[-1], matrix[:-1]).ravel()
        except ValueError:
            semantic = np.zeros(len(texts), dtype=float)

        total_weight = max(
            settings.RETRIEVAL_SEMANTIC_WEIGHT + settings.RETRIEVAL_LEXICAL_WEIGHT,
            1e-9,
        )
        combined = (
            settings.RETRIEVAL_SEMANTIC_WEIGHT * _normalize(semantic)
            + settings.RETRIEVAL_LEXICAL_WEIGHT * _normalize(lexical)
        ) / total_weight
        order = np.argsort(-combined)[:top_k]
        results: list[TrechoRecuperado] = []
        for index_position in order:
            score = float(combined[index_position])
            if score < settings.RETRIEVAL_MIN_RELEVANCE:
                continue
            chunk = chunks[int(index_position)]
            metadata = documents.get(chunk["document_id"], {})
            results.append(
                TrechoRecuperado(
                    chunk_id=chunk["id"],
                    document_id=chunk["document_id"],
                    trecho=chunk["text"],
                    fonte=chunk["filename"],
                    classificacao=chunk["classificacao"],
                    data_documento=_date_from_metadata(metadata.get("data_documento")),
                    versao_documento=metadata.get("version"),
                    relevancia=round(score, 4),
                )
            )
        return results
