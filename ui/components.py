from __future__ import annotations

import io
import json

import pandas as pd
import streamlit as st
from docx import Document


CONFIDENCE_LABELS = {
    "alta": "Alta",
    "media": "Média",
    "baixa": "Baixa",
}

QUESTION_TYPE_LABELS = {
    "contexto": "Contexto",
    "operacao": "Operação",
    "impacto": "Impacto",
    "maturidade": "Maturidade",
    "prioridade": "Prioridade",
    "decisao": "Decisão",
    "aderencia": "Aderência",
}


def _classification_label(value: str) -> str:
    labels = {
        "publico": "Público",
        "interno_nao_sensivel": "Interno não sensível",
        "restrito": "Restrito",
        "sensivel_confidencial": "Sensível ou confidencial",
    }
    return labels.get(value, value.replace("_", " ").title() if value else "—")


def render_evidence(items: list[dict]) -> None:
    st.subheader("Evidências recuperadas")
    st.caption(
        "Trechos selecionados pelo mecanismo de recuperação. Confirme origem, versão, atualidade "
        "e relação com a oportunidade antes de aceitar qualquer hipótese."
    )
    if not items:
        st.info("Nenhuma evidência foi recuperada. Revise a base e o contexto informado.")
        return

    relevant_items = [item for item in items if float(item.get("relevancia", 0) or 0) >= 0.5]
    unique_sources = {item.get("document_id") or item.get("fonte") for item in items}
    metric_one, metric_two, metric_three, metric_four = st.columns(4)
    metric_one.metric("Trechos", len(items))
    metric_two.metric("Fontes", len(unique_sources))
    metric_three.metric("Relevância ≥ 0,50", len(relevant_items))
    metric_four.metric(
        "Maior relevância",
        f"{max(float(item.get('relevancia', 0) or 0) for item in items):.2f}",
    )

    rows = [
        {
            "Fonte": item.get("fonte", ""),
            "Versão": item.get("versao_documento") or "—",
            "Classificação": _classification_label(item.get("classificacao", "")),
            "Data": item.get("data_documento") or "—",
            "Relevância": round(float(item.get("relevancia", 0) or 0), 3),
            "Prévia do trecho": (
                str(item.get("trecho", ""))[:180] + "…"
                if len(str(item.get("trecho", ""))) > 180
                else str(item.get("trecho", ""))
            ),
        }
        for item in items
    ]
    st.dataframe(
        pd.DataFrame(rows),
        width="stretch",
        hide_index=True,
        column_config={
            "Fonte": st.column_config.TextColumn(width="medium"),
            "Versão": st.column_config.TextColumn(width="small"),
            "Classificação": st.column_config.TextColumn(width="small"),
            "Data": st.column_config.TextColumn(width="small"),
            "Relevância": st.column_config.NumberColumn(format="%.3f", width="small"),
            "Prévia do trecho": st.column_config.TextColumn(width="large"),
        },
    )

    with st.expander("Abrir trechos completos", expanded=False):
        for index, item in enumerate(items, 1):
            source = item.get("fonte", "Fonte não identificada")
            relevance = float(item.get("relevancia", 0) or 0)
            st.markdown(f"**{index}. {source}** · relevância `{relevance:.3f}`")
            metadata = []
            if item.get("versao_documento"):
                metadata.append(f"versão {item['versao_documento']}")
            if item.get("classificacao"):
                metadata.append(_classification_label(item["classificacao"]))
            if item.get("data_documento"):
                metadata.append(f"data {item['data_documento']}")
            if item.get("chunk_id"):
                metadata.append(f"referência técnica `{item['chunk_id']}`")
            if metadata:
                st.caption(" · ".join(metadata))
            st.write(item.get("trecho", ""))
            if index < len(items):
                st.divider()


def render_hypotheses(data: dict) -> None:
    st.subheader("Matriz de hipóteses")
    st.caption(
        "Possibilidades formuladas a partir das evidências. Elas ainda precisam ser confirmadas "
        "na conversa com o potencial cliente."
    )
    hypotheses = data.get("hipoteses", []) if data else []
    if not hypotheses:
        st.info("Nenhuma hipótese foi gerada.")
        return

    confidence_counts = {
        level: sum(1 for item in hypotheses if item.get("confianca") == level)
        for level in ("alta", "media", "baixa")
    }
    columns = st.columns(3)
    for column, level in zip(columns, ("alta", "media", "baixa")):
        column.metric(f"Confiança {CONFIDENCE_LABELS[level].lower()}", confidence_counts[level])

    for hypothesis in hypotheses:
        confidence = hypothesis.get("confianca", "")
        dimension = str(hypothesis.get("dimensao", "")).replace("_", " ").title()
        with st.expander(
            f"{dimension} · confiança {CONFIDENCE_LABELS.get(confidence, confidence)}",
            expanded=confidence == "alta",
        ):
            st.markdown(f"**Hipótese:** {hypothesis.get('hipotese', '')}")
            rationale = hypothesis.get("fundamentacao_confianca")
            if rationale:
                st.markdown(f"**Fundamentação da confiança:** {rationale}")
            evidence = hypothesis.get("evidencias", [])
            if evidence:
                st.markdown("**Evidências relacionadas:**")
                for item in evidence:
                    if isinstance(item, dict):
                        description = item.get("descricao", "")
                        source = item.get("fonte", "Fonte não identificada")
                        chunk = f" · referência `{item['chunk_id']}`" if item.get("chunk_id") else ""
                        st.markdown(f"- {description} — `{source}`{chunk}")
                    else:
                        st.markdown(f"- {item}")
            else:
                st.warning("A hipótese não possui evidência vinculada e exige revisão cuidadosa.")
            st.markdown(
                f"**Pergunta de confirmação:** {hypothesis.get('pergunta_confirmacao', '')}"
            )


def questions_dataframe(data: dict) -> pd.DataFrame:
    questions = data.get("perguntas", []) if data else []
    return pd.DataFrame(
        [
            {
                "Tipo": q.get("tipo", ""),
                "Pergunta": q.get("pergunta", ""),
                "Prioridade": q.get("prioridade", 3),
                "Finalidade": q.get("finalidade", ""),
            }
            for q in questions
        ]
    )


def render_questions(data: dict) -> None:
    st.subheader("Roteiro diagnóstico")
    st.caption(
        "Prioridade 1 = indispensável; 2 = necessária para decidir o próximo passo; "
        "3 = útil; 4 = complementar; 5 = opcional."
    )
    questions = data.get("perguntas", []) if data else []
    if not questions:
        st.info("Nenhuma pergunta diagnóstica foi gerada.")
        return

    ordered = sorted(
        questions,
        key=lambda item: (int(item.get("prioridade", 3)), str(item.get("tipo", ""))),
    )
    for index, question in enumerate(ordered, 1):
        priority = int(question.get("prioridade", 3) or 3)
        question_type = QUESTION_TYPE_LABELS.get(
            question.get("tipo", ""),
            str(question.get("tipo", "")).replace("_", " ").title(),
        )
        with st.container(border=True):
            heading, badge = st.columns([6, 1])
            with heading:
                st.markdown(f"**{index}. {question.get('pergunta', '')}**")
            with badge:
                st.markdown(
                    f'<span class="priority-badge priority-{priority}">P{priority}</span>',
                    unsafe_allow_html=True,
                )
            st.caption(f"Tipo: {question_type}")
            st.markdown(f"**Finalidade:** {question.get('finalidade', '')}")


def render_qualification(data: dict) -> None:
    st.subheader("Qualificação da oportunidade")
    st.caption(
        "Leitura preliminar para orientar o próximo passo. A recomendação não substitui a decisão profissional."
    )
    if not data:
        st.info("A qualificação ainda não foi gerada.")
        return

    recommendation = str(data.get("recomendacao", "")).replace("_", " ").upper()
    st.markdown(
        f'<div class="recommendation-card"><span>Recomendação</span><strong>{recommendation or "—"}</strong></div>',
        unsafe_allow_html=True,
    )
    st.markdown(f"**Justificativa:** {data.get('justificativa', '')}")

    criteria = data.get("criterios", {})
    if criteria:
        st.dataframe(
            pd.DataFrame(
                [
                    {"Critério": key.replace("_", " ").title(), "Análise": value}
                    for key, value in criteria.items()
                ]
            ),
            hide_index=True,
            width="stretch",
            column_config={
                "Critério": st.column_config.TextColumn(width="small"),
                "Análise": st.column_config.TextColumn(width="large"),
            },
        )

    gaps = data.get("lacunas", [])
    if gaps:
        st.markdown("**Lacunas que ainda precisam ser confirmadas:**")
        for gap in gaps:
            st.markdown(f"- {gap}")


def _approved_script_docx(data: dict) -> bytes:
    document = Document()
    document.add_heading("PROSPECT-LLM — Roteiro aprovado", level=1)
    document.add_paragraph(f"Potencial cliente: {data.get('prospect', '—')}")
    document.add_paragraph(f"Aprovado por: {data.get('aprovado_por', '—')}")
    document.add_paragraph(f"Data de aprovação: {data.get('data_aprovacao', '—')}")
    if data.get("recomendacao_final"):
        document.add_paragraph(
            "Recomendação: " + str(data["recomendacao_final"]).replace("_", " ").upper()
        )
    document.add_heading("Contexto factual", level=2)
    document.add_paragraph(data.get("resumo_factual", ""))
    document.add_heading("Perguntas finais", level=2)
    for question in sorted(
        data.get("perguntas_finais", []),
        key=lambda item: int(item.get("prioridade", 3)),
    ):
        document.add_paragraph(
            f"P{question.get('prioridade', 3)} — {question.get('pergunta', '')}",
            style="List Bullet",
        )
        document.add_paragraph(f"Finalidade: {question.get('finalidade', '')}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_approved_script(data: dict) -> None:
    st.subheader("Roteiro aprovado")
    if not data:
        st.info("O conteúdo do roteiro aprovado não foi encontrado.")
        return

    top_left, top_right = st.columns([3, 1])
    with top_left:
        st.markdown(f"### {data.get('prospect', 'Potencial cliente')}")
        st.caption(f"Aprovado por: {data.get('aprovado_por', '—')}")
    with top_right:
        st.caption("Data de aprovação")
        st.markdown(f"**{data.get('data_aprovacao', '—')}**")

    recommendation = data.get("recomendacao_final")
    if recommendation:
        st.markdown(
            f"**Recomendação registrada:** {str(recommendation).replace('_', ' ').upper()}"
        )

    with st.expander("Contexto factual registrado", expanded=False):
        st.text(data.get("resumo_factual", ""))

    render_questions({"perguntas": data.get("perguntas_finais", [])})

    json_bytes = json.dumps(data, ensure_ascii=False, indent=2, default=str).encode("utf-8")
    csv_bytes = questions_dataframe({"perguntas": data.get("perguntas_finais", [])}).to_csv(
        index=False
    ).encode("utf-8-sig")
    docx_bytes = _approved_script_docx(data)
    one, two, three = st.columns(3)
    one.download_button(
        "Baixar JSON",
        data=json_bytes,
        file_name="roteiro_aprovado.json",
        mime="application/json",
        width="stretch",
    )
    two.download_button(
        "Baixar CSV",
        data=csv_bytes,
        file_name="roteiro_aprovado.csv",
        mime="text/csv",
        width="stretch",
    )
    three.download_button(
        "Baixar DOCX",
        data=docx_bytes,
        file_name="roteiro_aprovado.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        width="stretch",
    )
